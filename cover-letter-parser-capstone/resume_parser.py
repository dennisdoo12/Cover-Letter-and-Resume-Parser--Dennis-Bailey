#If the spaCy model isn't installed, the parser automatically falls back to
#regex-only extraction, so it still runs.

import re
import os

try:
    import spacy
    try:
        _NLP = spacy.load("en_core_web_sm")
    except OSError:
        _NLP = None            # library present, model not downloaded
except ImportError:
    _NLP = None                # spaCy not installed at all


# Text extraction  (PDF and DOCX -> raw text)
def extract_text(file):
    #Return the raw text of a resume file (.pdf or .docx).
    #Returns an empty string if no text could be extracted
    ext = os.path.splitext(file)[1].lower()
    if ext == ".pdf":
        return _extract_pdf(file)
    elif ext in (".docx", ".doc"):
        return _extract_docx(file)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _extract_pdf(file):
    import fitz  # PyMuPDF
    text_parts = []
    with fitz.open(file) as doc:
        for page in doc:
            text_parts.append(page.get_text("text"))
    return "\n".join(text_parts).strip()


def _extract_docx(file):
    import docx
    document = docx.Document(file)
    parts = [p.text for p in document.paragraphs]
    # pull text out of any tables
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts).strip()


EMAIL_RE = re.compile(r"[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}")
# Matches formats like 123-456-7890, (123) 456-7890, +1 123 456 7890
PHONE_RE = re.compile(r"(?:\+?\d{1,2}[\s.\-]?)?(?:\(?\d{3}\)?[\s.\-]?)\d{3}[\s.\-]?\d{4}")
YEAR_RE = re.compile(r"(19|20)\d{2}")


def extract_email(text):
    match = EMAIL_RE.search(text)
    return match.group(0) if match else None


def extract_phone(text):
    match = PHONE_RE.search(text)
    return match.group(0).strip() if match else None


def extract_name(text):
    #Uses spaCy PERSON entities when the model is available (more reliable);
    #otherwise falls back to a heuristic approach: the first non-empty line that looks like a name (2-3 capitalized words, no digits/@).
    if _NLP is not None:
        doc = _NLP(text[:1000])  # name is near the top
        for ent in doc.ents:
            if ent.label_ == "PERSON":
                return ent.text.strip()
    #fallback
    for line in text.splitlines():
        line = line.strip()
        if not line or "@" in line or any(ch.isdigit() for ch in line):
            continue
        words = line.split()
        if 2 <= len(words) <= 3 and all(w[0].isupper() for w in words if w):
            return line
    return None


# skills dictionary. Keys are the CANONICAL skill name; the list
# holds aliases/variants that should all normalize to that canonical name.
# plans to extend list
SKILLS = {
    "Python":        ["python"],
    "JavaScript":    ["javascript", "js", "ecmascript"],
    "TypeScript":    ["typescript", "ts"],
    "Java":          ["java"],
    "C++":           ["c++", "cpp"],
    "C#":            ["c#", "csharp"],
    "SQL":           ["sql", "mysql", "postgresql", "postgres", "sqlite"],
    "React":         ["react", "react.js", "reactjs"],
    "Node.js":       ["node", "node.js", "nodejs"],
    "HTML":          ["html", "html5"],
    "CSS":           ["css", "css3"],
    "Git":           ["git", "github", "version control"],
    "Docker":        ["docker", "containerization"],
    "AWS":           ["aws", "amazon web services"],
    "Machine Learning": ["machine learning", "ml"],
    "REST API":      ["rest", "rest api", "restful", "api"],
    "Django":        ["django"],
    "Flask":         ["flask"],
    "spaCy":         ["spacy"],
    "MongoDB":       ["mongodb", "mongo", "nosql"],
}

# longest aliases first so multi-word aliases match before single words.
_ALIAS_TO_CANON = []
for canon, aliases in SKILLS.items():
    for alias in aliases:
        _ALIAS_TO_CANON.append((alias, canon))

def _alias_length(pair):
    return len(pair[0])

_ALIAS_TO_CANON.sort(key=_alias_length, reverse=True)

def extract_skills(text):
    #Return a sorted list of canonical skills found in the text.
    #Matching is done on word boundaries against the alias table, then normalized names
    lowered = text.lower()
    found = set()
    for alias, canon in _ALIAS_TO_CANON:
        # word-boundary match; re.escape handles things like c++ and c#
        # '.' is in the lookarounds so "js" doesn't match inside "Node.js"
        pattern = r"(?<![\w+#.])" + re.escape(alias) + r"(?![\w+#.])"
        if re.search(pattern, lowered):
            found.add(canon)
    return sorted(found)


# Experience
def extract_experience_years(text):
    # explicit statement, e.g. "5 years of experience"
    stated = re.search(r"(\d{1,2})\+?\s*years?\s+(?:of\s+)?experience", text, re.I)
    if stated:
        return int(stated.group(1))

    # estimate by summing individual job date ranges, e.g. "2020 - 2024"
    # or "2022 - Present". Summing (not global min/max) avoids counting the
    # gap between a graduation year and a first job as work experience.
    CURRENT_YEAR = 2026
    range_re = re.compile(
        r"((?:19|20)\d{2})\s*[-\u2013to]+\s*((?:19|20)\d{2}|present|current|now)",
        re.I,
    )
    total = 0
    for start, end in range_re.findall(text):
        start = int(start)
        end = CURRENT_YEAR if not end.isdigit() else int(end)
        span = end - start
        if 0 < span <= 50:          # realistic check
            total += span
    return total if total else None


# parse a whole resume into a structured dict
def parse_resume(file):
    raw_text = extract_text(file)

    if not raw_text:
        return {
            "parsed": False,
            "reason": "No extractable text (possibly a scanned/image-only file).",
            "file": os.path.basename(file),
        }

    return {
        "parsed": True,
        "file": os.path.basename(file),
        "name": extract_name(raw_text),
        "email": extract_email(raw_text),
        "phone": extract_phone(raw_text),
        "skills": extract_skills(raw_text),
        "experience_years": extract_experience_years(raw_text),
    }


# manual test
if __name__ == "__main__":
    import json
    import sys

    if len(sys.argv) > 1:
        # parse a real file passed on the command line
        result = parse_resume(sys.argv[1])
    else:
        sample = """
        Jane Doe
        jane.doe@example.com | (415) 555-0198

        SUMMARY
        Backend developer with 4 years of experience building web services.

        EXPERIENCE
        Software Engineer, Acme Corp        2021 - 2024
        - Built REST APIs in Python and Node.js
        - Worked with SQL databases and Docker

        Junior Developer, StartupXYZ        2019 - 2021
        - Front-end work in React, JavaScript, HTML and CSS

        SKILLS: Python, JS, React, SQL, Git, AWS
        """
        # run the field extractors directly on the sample text
        result = {
            "parsed": True,
            "name": extract_name(sample),
            "email": extract_email(sample),
            "phone": extract_phone(sample),
            "skills": extract_skills(sample),
            "experience_years": extract_experience_years(sample),
        }

    print(json.dumps(result, indent=2))
    print("\nspaCy model active:", _NLP is not None)
