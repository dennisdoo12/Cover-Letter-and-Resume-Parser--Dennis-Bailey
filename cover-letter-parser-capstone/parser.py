import re
import sqlite3
from pathlib import Path

from docx import Document
from pypdf import PdfReader


# ==================================================
# DATABASE
# ==================================================

BASE_DIR = Path(__file__).resolve().parent
ONET_DB = BASE_DIR / "onet.db"


# ==================================================
# GENERAL CONFIGURATION
# ==================================================

ACTION_VERBS = [
    "achieved",
    "administered",
    "analyzed",
    "automated",
    "built",
    "collaborated",
    "configured",
    "coordinated",
    "created",
    "delivered",
    "deployed",
    "designed",
    "developed",
    "diagnosed",
    "directed",
    "engineered",
    "generated",
    "implemented",
    "improved",
    "increased",
    "installed",
    "launched",
    "led",
    "maintained",
    "managed",
    "migrated",
    "monitored",
    "optimized",
    "organized",
    "produced",
    "reduced",
    "resolved",
    "supported",
    "tested",
    "trained",
    "troubleshot",
    "upgraded"
]


STOP_WORDS = {
    "the",
    "and",
    "for",
    "with",
    "that",
    "this",
    "from",
    "your",
    "our",
    "you",
    "are",
    "will",
    "job",
    "role",
    "position",
    "candidate",
    "experience",
    "skills",
    "work",
    "team",
    "company",
    "have",
    "has",
    "who",
    "but",
    "not",
    "all",
    "can",
    "their",
    "they",
    "its",
    "into",
    "about",
    "years",
    "year",
    "required",
    "preferred",
    "responsibilities",
    "requirements",
    "including",
    "using"
}


# ==================================================
# DATABASE HELPERS
# ==================================================

def get_db_connection():
    if not ONET_DB.exists():
        raise FileNotFoundError(
            f"O*NET database was not found at: {ONET_DB}"
        )

    connection = sqlite3.connect(ONET_DB)

    connection.row_factory = sqlite3.Row

    return connection


def get_onet_software_skills():
    """
    Load every unique software / technology example from O*NET.

    Expected table:
        software_skills

    Expected columns created by build_database.py:
        workplace_example
        hot_technology
        in_demand
    """

    connection = get_db_connection()

    try:
        rows = connection.execute(
            """
            SELECT
                workplace_example,
                MAX(hot_technology) AS hot_technology,
                MAX(in_demand) AS in_demand
            FROM software_skills
            WHERE workplace_example IS NOT NULL
              AND TRIM(workplace_example) != ''
            GROUP BY LOWER(workplace_example)
            """
        ).fetchall()

        return [
            {
                "name": row["workplace_example"],
                "hot_technology": (
                    str(row["hot_technology"]).upper() == "Y"
                ),
                "in_demand": (
                    str(row["in_demand"]).upper() == "Y"
                )
            }
            for row in rows
        ]

    finally:
        connection.close()


def get_onet_essential_skill_names():
    """
    Load unique Essential Skill names from O*NET.
    """

    connection = get_db_connection()

    try:
        rows = connection.execute(
            """
            SELECT DISTINCT element_name
            FROM essential_skills
            WHERE element_name IS NOT NULL
              AND TRIM(element_name) != ''
            ORDER BY element_name
            """
        ).fetchall()

        return [
            row["element_name"]
            for row in rows
        ]

    finally:
        connection.close()


# ==================================================
# DOCUMENT READING
# ==================================================

def read_document(path):

    suffix = Path(path).suffix.lower()

    if suffix == ".txt":

        return Path(path).read_text(
            encoding="utf-8",
            errors="ignore"
        )

    if suffix == ".pdf":

        return "\n".join(
            (page.extract_text() or "").strip()
            for page in PdfReader(path).pages
        )

    if suffix == ".docx":

        doc = Document(path)

        parts = [
            paragraph.text.strip()
            for paragraph in doc.paragraphs
            if paragraph.text.strip()
        ]

        for table in doc.tables:

            for row in table.rows:

                for cell in row.cells:

                    if cell.text.strip():

                        parts.append(
                            cell.text.strip()
                        )

        return "\n".join(parts)

    raise ValueError(
        "Unsupported file type."
    )


# ==================================================
# BASIC EXTRACTION
# ==================================================

def extract(pattern, text):

    match = re.search(
        pattern,
        text,
        re.IGNORECASE
    )

    return (
        match.group(0).strip()
        if match
        else None
    )


def extract_name(text):

    ignored = {
        "dear hiring manager",
        "dear recruiter",
        "to whom it may concern",
        "professional summary",
        "summary",
        "resume",
        "curriculum vitae"
    }

    lines = [
        line.strip()
        for line in text.splitlines()
        if line.strip()
    ][:10]

    for line in lines:

        low = (
            line.lower()
            .strip(" ,:")
        )

        if (
            low in ignored
            or "@" in line
            or re.search(r"\d", line)
        ):
            continue

        words = line.split()

        if (
            2 <= len(words) <= 4
            and all(
                re.fullmatch(
                    r"[A-Za-z.'-]+",
                    word
                )
                for word in words
            )
        ):

            return line

    return None


def phrase_in_text(phrase, text):
    """
    Match an O*NET phrase against document text.

    Examples:
        Python
        Microsoft Excel
        Amazon Web Services AWS
    """

    phrase = phrase.strip()

    if not phrase:
        return False

    pattern = (
        r"(?<!\w)"
        + re.escape(phrase)
        + r"(?!\w)"
    )

    return bool(
        re.search(
            pattern,
            text,
            re.IGNORECASE
        )
    )


# ==================================================
# O*NET SKILL EXTRACTION
# ==================================================

def extract_onet_skills(text):

    found = []

    seen = set()

    for skill in get_onet_software_skills():

        name = skill["name"]

        normalized = name.lower()

        if normalized in seen:
            continue

        if phrase_in_text(
            name,
            text
        ):

            found.append(skill)

            seen.add(normalized)

    return found


def extract_essential_skills(text):

    found = []

    for skill in get_onet_essential_skill_names():

        if phrase_in_text(
            skill,
            text
        ):

            found.append(skill)

    return found


def extract_skills(text):

    """
    Maintain the old front-end format.

    Returns a simple list of names.
    """

    software = extract_onet_skills(text)

    essential = extract_essential_skills(text)

    names = [
        skill["name"]
        for skill in software
    ]

    names.extend(
        essential
    )

    # Remove duplicates while preserving order.
    return list(
        dict.fromkeys(names)
    )


def count_words(text):

    return len(
        re.findall(
            r"\b[\w+#.-]+\b",
            text
        )
    )


def summarize(
    text,
    max_chars=420
):

    clean = re.sub(
        r"\s+",
        " ",
        text
    ).strip()

    if len(clean) <= max_chars:
        return clean

    return (
        clean[:max_chars]
        .rsplit(" ", 1)[0]
        + "..."
    )


# ==================================================
# O*NET JOB MATCHING
# ==================================================

def job_match(
    text,
    job_description
):

    if not job_description.strip():

        return {
            "matched_keywords": [],
            "missing_keywords": [],
            "hot_technology_matches": [],
            "in_demand_matches": [],
            "match_percentage": None
        }

    # ----------------------------------------------
    # Find O*NET technologies mentioned in job ad
    # ----------------------------------------------

    job_software_skills = []

    for skill in get_onet_software_skills():

        if phrase_in_text(
            skill["name"],
            job_description
        ):

            job_software_skills.append(
                skill
            )

    # ----------------------------------------------
    # Find O*NET essential skills mentioned in job ad
    # ----------------------------------------------

    job_essential_skills = []

    for skill in get_onet_essential_skill_names():

        if phrase_in_text(
            skill,
            job_description
        ):

            job_essential_skills.append(
                skill
            )

    # ----------------------------------------------
    # Compare against resume / cover letter
    # ----------------------------------------------

    matched = []

    missing = []

    hot_matches = []

    demand_matches = []

    matched_weight = 0

    total_weight = 0


    for skill in job_software_skills:

        name = skill["name"]

        # Base O*NET technology = 2 points
        weight = 2

        # Hot Technology receives additional weight
        if skill["hot_technology"]:
            weight += 1

        # In-demand technology receives additional weight
        if skill["in_demand"]:
            weight += 1

        total_weight += weight

        if phrase_in_text(
            name,
            text
        ):

            matched.append(name)

            matched_weight += weight

            if skill["hot_technology"]:

                hot_matches.append(
                    name
                )

            if skill["in_demand"]:

                demand_matches.append(
                    name
                )

        else:

            missing.append(name)


    # Essential skills are lower-weighted than software
    # technologies in this prototype.

    for skill in job_essential_skills:

        weight = 1

        total_weight += weight

        if phrase_in_text(
            skill,
            text
        ):

            matched.append(skill)

            matched_weight += weight

        else:

            missing.append(skill)


    # ----------------------------------------------
    # FALLBACK
    # ----------------------------------------------
    #
    # Some job descriptions may contain no exact
    # O*NET skill phrases.
    #
    # In that case, perform the old general keyword
    # comparison rather than returning 0%.
    # ----------------------------------------------

    if total_weight == 0:

        job_words = {
            word
            for word in re.findall(
                r"\b[a-zA-Z]"
                r"[a-zA-Z+#.-]{2,}\b",
                job_description.lower()
            )
            if word not in STOP_WORDS
        }

        text_words = set(
            re.findall(
                r"\b[a-zA-Z]"
                r"[a-zA-Z+#.-]{2,}\b",
                text.lower()
            )
        )

        general_matched = sorted(
            job_words
            & text_words
        )

        general_missing = sorted(
            job_words
            - text_words
        )

        percentage = (
            round(
                len(general_matched)
                / len(job_words)
                * 100
            )
            if job_words
            else 0
        )

        return {
            "matched_keywords":
                general_matched[:30],

            "missing_keywords":
                general_missing[:30],

            "hot_technology_matches": [],

            "in_demand_matches": [],

            "match_percentage":
                percentage,

            "scoring_method":
                "general_keyword_fallback"
        }


    percentage = round(
        matched_weight
        / total_weight
        * 100
    )


    return {

        "matched_keywords":
            list(
                dict.fromkeys(
                    matched
                )
            )[:40],

        "missing_keywords":
            list(
                dict.fromkeys(
                    missing
                )
            )[:40],

        "hot_technology_matches":
            list(
                dict.fromkeys(
                    hot_matches
                )
            ),

        "in_demand_matches":
            list(
                dict.fromkeys(
                    demand_matches
                )
            ),

        "match_percentage":
            percentage,

        "matched_weight":
            matched_weight,

        "total_weight":
            total_weight,

        "scoring_method":
            "onet_weighted"
    }


# ==================================================
# SHARED INFORMATION
# ==================================================

def extract_common_details(text):

    email = extract(
        r"\b[A-Za-z0-9._%+-]+"
        r"@[A-Za-z0-9.-]+"
        r"\.[A-Za-z]{2,}\b",
        text
    )

    phone = extract(
        r"(?:\+?1[-.\s]?)?"
        r"(?:\(?\d{3}\)?[-.\s]?)"
        r"\d{3}[-.\s]?\d{4}",
        text
    )

    linkedin = extract(
        r"(?:https?://)?"
        r"(?:www\.)?"
        r"linkedin\.com/in/"
        r"[A-Za-z0-9_-]+/?",
        text
    )

    name = extract_name(
        text
    )

    skills = extract_skills(
        text
    )

    actions = [
        verb
        for verb in ACTION_VERBS
        if re.search(
            rf"\b{verb}\b",
            text,
            re.I
        )
    ]

    metrics = re.findall(
        r"\b\d+(?:\.\d+)?%\b"
        r"|\$\s?\d[\d,]*(?:\.\d+)?"
        r"|\b\d[\d,]*\+\b",
        text
    )

    return {

        "name":
            name,

        "email":
            email,

        "phone":
            phone,

        "linkedin":
            linkedin,

        "skills":
            skills,

        "action_verbs":
            actions,

        "metrics":
            metrics
    }


# ==================================================
# RATING HELPER
# ==================================================

def get_rating(total):

    if total >= 85:
        return "Excellent"

    if total >= 70:
        return "Strong"

    if total >= 55:
        return "Average"

    if total >= 40:
        return "Needs Improvement"

    return "Weak"


# ==================================================
# COVER LETTER PARSER
# ==================================================

def parse_cover_letter(
    text,
    job_description=""
):

    details = extract_common_details(
        text
    )

    name = details["name"]

    email = details["email"]

    phone = details["phone"]

    linkedin = details["linkedin"]

    skills = details["skills"]

    actions = details["action_verbs"]

    metrics = details["metrics"]

    words = count_words(
        text
    )

    sentences = max(
        1,
        len([
            sentence
            for sentence in re.split(
                r"[.!?]+",
                text
            )
            if sentence.strip()
        ])
    )

    avg_sentence = (
        words
        / sentences
    )


    greeting = bool(
        re.search(
            r"(?mi)^\s*"
            r"(dear .+|to whom it may concern)"
            r"[,:\s]*$",
            text
        )
    )


    closing = bool(
        re.search(
            r"(?mi)^\s*"
            r"(sincerely"
            r"|regards"
            r"|best regards"
            r"|thank you"
            r"|respectfully)"
            r"[,:\s]*$",
            text
        )
    )


    interest = bool(
        re.search(
            r"\b("
            r"apply(?:ing)? for"
            r"|interested in"
            r"|excited to apply"
            r"|position"
            r"|role"
            r")\b",
            text,
            re.I
        )
    )


    company = bool(
        re.search(
            r"\bat "
            r"[A-Z]"
            r"[A-Za-z0-9&.' -]{2,}\b",
            text
        )
    )


    match = job_match(
        text,
        job_description
    )


    scores = {}


    scores[
        "Contact Information"
    ] = (

        (3 if name else 0)

        + (3 if email else 0)

        + (2 if phone else 0)

        + (2 if linkedin else 0)

    )


    scores[
        "Professional Structure"
    ] = (

        (5 if greeting else 0)

        + (5 if closing else 0)

        + (5 if sentences >= 3 else 0)

    )


    relevance = (
        (8 if interest else 0)
        + (6 if company else 0)
    )


    if (
        match["match_percentage"]
        is not None
    ):

        if (
            match["match_percentage"]
            >= 70
        ):

            relevance += 6

        elif (
            match["match_percentage"]
            >= 45
        ):

            relevance += 5

        elif (
            match["match_percentage"]
            >= 25
        ):

            relevance += 3

        elif (
            match["match_percentage"]
            > 0
        ):

            relevance += 1

    elif skills:

        relevance += 4


    scores[
        "Job Relevance"
    ] = min(
        relevance,
        20
    )


    skill_score = (

        min(
            len(skills) * 2,
            12
        )

        +

        (
            4
            if len(actions) >= 4

            else 2
            if len(actions) >= 2

            else 0
        )

        +

        (
            4
            if len(metrics) >= 2

            else 2
            if len(metrics) == 1

            else 0
        )

    )


    scores[
        "Skills and Qualifications"
    ] = min(
        skill_score,
        20
    )


    clarity = (

        (
            6
            if 180 <= words <= 500

            else 4
            if 120 <= words <= 650

            else 0
        )

        +

        (
            5
            if avg_sentence <= 24

            else 3
            if avg_sentence <= 30

            else 0
        )

    )


    if not re.search(
        r"\b("
        r"stuff"
        r"|whatever"
        r"|probably"
        r"|kinda"
        r"|sort of"
        r")\b",
        text,
        re.I
    ):

        clarity += 4


    scores[
        "Writing Clarity"
    ] = min(
        clarity,
        15
    )


    scores[
        "Personalization"
    ] = min(

        (5 if company else 0)

        +

        (
            3
            if re.search(
                r"\b("
                r"because"
                r"|particularly"
                r"|specifically"
                r")\b",
                text,
                re.I
            )

            else 0
        )

        +

        (
            2
            if "your company"
            not in text.lower()

            else 0
        ),

        10
    )


    scores[
        "Completeness"
    ] = min(

        (4 if words >= 150 else 0)

        + (2 if greeting else 0)

        + (2 if closing else 0)

        + (2 if interest else 0),

        10
    )


    total = max(
        1,
        min(
            sum(
                scores.values()
            ),
            100
        )
    )


    rating = get_rating(
        total
    )


    strengths = []

    improvements = []


    checks = [

        (
            scores[
                "Contact Information"
            ] >= 8,

            "Includes clear applicant contact information.",

            "Add your full name, email, phone number, and LinkedIn profile."
        ),

        (
            scores[
                "Professional Structure"
            ] >= 12,

            "Uses a clear professional cover letter structure.",

            "Include a professional greeting, body paragraphs, and closing."
        ),

        (
            scores[
                "Job Relevance"
            ] >= 15,

            "Connects the applicant's background to the target role.",

            "Explain why you want this role and connect experience to requirements."
        ),

        (
            scores[
                "Skills and Qualifications"
            ] >= 14,

            "Provides strong evidence of relevant skills and qualifications.",

            "Add more relevant O*NET skills, accomplishments, and measurable results."
        ),

        (
            scores[
                "Writing Clarity"
            ] >= 12,

            "Writing is concise, professional, and easy to read.",

            "Use clearer wording, shorter sentences, and a professional tone."
        ),

        (
            scores[
                "Personalization"
            ] >= 7,

            "The letter appears tailored to the employer.",

            "Mention the employer by name and explain why it interests you."
        ),

        (
            scores[
                "Completeness"
            ] >= 8,

            "Includes the major expected cover letter components.",

            "Add a clear introduction, qualifications, employer connection, and closing."
        )
    ]


    for passed, good, fix in checks:

        if passed:

            strengths.append(
                good
            )

        else:

            improvements.append(
                fix
            )


    return {

        **details,

        "word_count":
            words,

        "summary":
            summarize(text),

        "score":
            total,

        "rating":
            rating,

        "category_scores":
            scores,

        "strengths":
            strengths,

        "improvements":
            improvements,

        "readability": {

            "sentence_count":
                sentences,

            "average_sentence_length":
                round(
                    avg_sentence,
                    1
                )
        },

        "job_match":
            match,

        "skill_source":
            "O*NET 30.3"
    }


# ==================================================
# RESUME PARSER
# ==================================================

def parse_resume(
    text,
    job_description=""
):

    details = extract_common_details(
        text
    )


    skills = details[
        "skills"
    ]


    actions = details[
        "action_verbs"
    ]


    metrics = details[
        "metrics"
    ]


    words = count_words(
        text
    )


    match = job_match(
        text,
        job_description
    )


    sections = {

        "Summary": bool(
            re.search(
                r"(?mi)^\s*"
                r"(professional summary"
                r"|summary"
                r"|profile"
                r"|objective)"
                r"\s*:?\s*$",
                text
            )
        ),

        "Experience": bool(
            re.search(
                r"(?mi)^\s*"
                r"(work experience"
                r"|professional experience"
                r"|experience"
                r"|employment)"
                r"\s*:?\s*$",
                text
            )
        ),

        "Education": bool(
            re.search(
                r"(?mi)^\s*"
                r"education"
                r"\s*:?\s*$",
                text
            )
        ),

        "Skills": bool(
            re.search(
                r"(?mi)^\s*"
                r"(skills"
                r"|technical skills"
                r"|core competencies)"
                r"\s*:?\s*$",
                text
            )
        ),

        "Projects": bool(
            re.search(
                r"(?mi)^\s*"
                r"(projects"
                r"|academic projects"
                r"|personal projects)"
                r"\s*:?\s*$",
                text
            )
        ),

        "Certifications": bool(
            re.search(
                r"(?mi)^\s*"
                r"(certifications"
                r"|licenses"
                r"|certificates)"
                r"\s*:?\s*$",
                text
            )
        )
    }


    education_terms = bool(
        re.search(
            r"\b("
            r"B\.?A\.?"
            r"|B\.?S\.?"
            r"|M\.?A\.?"
            r"|M\.?S\.?"
            r"|Ph\.?D\.?"
            r"|bachelor'?s?"
            r"|master'?s?"
            r"|associate'?s?"
            r"|university"
            r"|college"
            r")\b",
            text,
            re.I
        )
    )


    bullet_count = len(
        re.findall(
            r"(?m)^\s*"
            r"(?:[-•*]|\d+[.)])"
            r"\s+",
            text
        )
    )


    scores = {}


    # ----------------------------------------------
    # CONTACT INFORMATION / 10
    # ----------------------------------------------

    scores[
        "Contact Information"
    ] = min(

        (
            3
            if details["name"]
            else 0
        )

        +

        (
            3
            if details["email"]
            else 0
        )

        +

        (
            2
            if details["phone"]
            else 0
        )

        +

        (
            2
            if details["linkedin"]
            else 0
        ),

        10
    )


    # ----------------------------------------------
    # RESUME STRUCTURE / 20
    # ----------------------------------------------

    structure_score = 0


    structure_score += (
        5
        if sections["Experience"]
        else 0
    )


    structure_score += (
        5
        if sections["Education"]
        else 0
    )


    structure_score += (
        5
        if sections["Skills"]
        else 0
    )


    structure_score += (
        3
        if (
            sections["Summary"]
            or sections["Projects"]
        )
        else 0
    )


    structure_score += (
        2
        if bullet_count >= 3
        else 0
    )


    scores[
        "Resume Structure"
    ] = min(
        structure_score,
        20
    )


    # ----------------------------------------------
    # SKILLS AND JOB RELEVANCE / 25
    # ----------------------------------------------

    relevance_score = min(
        len(skills) * 1.5,
        10
    )


    if (
        match["match_percentage"]
        is not None
    ):

        percentage = (
            match[
                "match_percentage"
            ]
        )


        if percentage >= 80:

            relevance_score += 15

        elif percentage >= 65:

            relevance_score += 13

        elif percentage >= 50:

            relevance_score += 11

        elif percentage >= 35:

            relevance_score += 8

        elif percentage >= 20:

            relevance_score += 5

        elif percentage > 0:

            relevance_score += 2


    else:

        relevance_score += min(
            len(skills),
            8
        )


    scores[
        "Skills and Job Relevance"
    ] = min(
        round(
            relevance_score
        ),
        25
    )


    # ----------------------------------------------
    # EXPERIENCE IMPACT / 20
    # ----------------------------------------------

    impact_score = (

        (
            8
            if len(actions) >= 8

            else 6
            if len(actions) >= 5

            else 4
            if len(actions) >= 3

            else 2
            if len(actions) >= 1

            else 0
        )

        +

        (
            8
            if len(metrics) >= 4

            else 6
            if len(metrics) >= 2

            else 3
            if len(metrics) == 1

            else 0
        )

        +

        (
            4
            if bullet_count >= 5

            else 2
            if bullet_count >= 2

            else 0
        )

    )


    scores[
        "Experience Impact"
    ] = min(
        impact_score,
        20
    )


    # ----------------------------------------------
    # EDUCATION / 10
    # ----------------------------------------------

    education_score = 0


    education_score += (
        6
        if sections["Education"]
        else 0
    )


    education_score += (
        4
        if education_terms
        else 0
    )


    scores[
        "Education"
    ] = min(
        education_score,
        10
    )


    # ----------------------------------------------
    # ATS READABILITY / 15
    # ----------------------------------------------

    ats_score = 0


    ats_score += (

        5
        if 150 <= words <= 1200

        else 3
        if 100 <= words <= 1600

        else 0

    )


    section_count = len([
        value
        for value
        in sections.values()
        if value
    ])


    ats_score += (
        4
        if section_count >= 3
        else 2
    )


    ats_score += (
        3
        if bullet_count >= 3
        else 1
    )


    has_dates = bool(
        re.search(
            r"\b(?:19|20)\d{2}\b",
            text
        )
    )


    ats_score += (
        3
        if has_dates
        else 1
    )


    scores[
        "ATS Readability"
    ] = min(
        ats_score,
        15
    )


    total = max(
        1,
        min(
            sum(
                scores.values()
            ),
            100
        )
    )


    rating = get_rating(
        total
    )


    strengths = []

    improvements = []


    checks = [

        (
            scores[
                "Contact Information"
            ] >= 8,

            "Includes clear applicant contact information.",

            "Add your name, email, phone number, and LinkedIn profile."
        ),

        (
            scores[
                "Resume Structure"
            ] >= 15,

            "Uses recognizable resume sections and organization.",

            "Use clear sections such as Experience, Education, and Skills."
        ),

        (
            scores[
                "Skills and Job Relevance"
            ] >= 18,

            "Shows strong alignment with O*NET skills found in the target job.",

            "Add O*NET-recognized skills that genuinely match the target job requirements."
        ),

        (
            scores[
                "Experience Impact"
            ] >= 14,

            "Uses action-oriented accomplishments and measurable results.",

            "Strengthen experience bullets with action verbs and measurable outcomes."
        ),

        (
            scores[
                "Education"
            ] >= 8,

            "Education information is clearly represented.",

            "Add a clearly labeled Education section with degree and school information."
        ),

        (
            scores[
                "ATS Readability"
            ] >= 11,

            "The resume has a structure that should be relatively easy to parse.",

            "Use standard section headings, concise bullets, and clear dates."
        )
    ]


    for passed, good, fix in checks:

        if passed:

            strengths.append(
                good
            )

        else:

            improvements.append(
                fix
            )


    return {

        **details,

        "word_count":
            words,

        "summary":
            summarize(text),

        "score":
            total,

        "rating":
            rating,

        "category_scores":
            scores,

        "strengths":
            strengths,

        "improvements":
            improvements,

        "sections_found":
            sections,

        "bullet_count":
            bullet_count,

        "job_match":
            match,

        "skill_source":
            "O*NET 30.3"
    }